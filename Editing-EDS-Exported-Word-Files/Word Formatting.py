import os
import zipfile
import tempfile
from pathlib import Path
from PIL import Image
from tkinter import Tk, filedialog
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# === Original EDX image sizes (cm) ===
BSE_ORIG = (12.16, 15.66)
SPEC_ORIG = (9.00, 15.37)
MAP_ORIG = (7.02, 7.21)
TOL = 0.2

# === Formatted target sizes (cm) ===
BSE_FMT = (5.34, 7.25)
SPEC_FMT = (4.75, 7.64)
MAP_FMT = (3.76, 3.66)

def extract_ordered_images(docx_file, extract_dir):
    extract_dir = Path(extract_dir)
    extract_dir.mkdir(parents=True, exist_ok=True)
    doc = Document(docx_file)
    rels = doc.part._rels
    img_order = []

    for p in doc.paragraphs:
        for run in p.runs:
            if "graphic" in run._element.xml and "r:embed=" in run._element.xml:
                rId = run._element.xpath(".//a:blip")[0].get(qn("r:embed"))
                if rId in rels:
                    target = rels[rId]._target
                    if hasattr(target, 'partname') and "media" in target.partname:
                        fname = Path(target.partname).name
                        img_order.append(fname)

    image_map = {}
    with zipfile.ZipFile(docx_file, 'r') as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                img_data = z.read(name)
                fname = name.split("/")[-1]
                out_path = extract_dir / fname
                with open(out_path, "wb") as f:
                    f.write(img_data)
                image_map[fname] = out_path

    ordered_paths = [image_map[f] for f in img_order if f in image_map]
    return ordered_paths

def size_cm(img_path):
    img = Image.open(img_path)
    dpi = img.info.get("dpi", (96, 96))[0]
    w, h = img.size
    return round(h / dpi * 2.54, 2), round(w / dpi * 2.54, 2)

def classify(img_path):
    h, w = size_cm(img_path)
    if abs(h - BSE_ORIG[0]) < TOL and abs(w - BSE_ORIG[1]) < TOL:
        return "BSE"
    elif abs(h - SPEC_ORIG[0]) < TOL and abs(w - SPEC_ORIG[1]) < TOL:
        return "SPEC"
    elif abs(h - MAP_ORIG[0]) < TOL and abs(w - MAP_ORIG[1]) < TOL:
        return "MAP"
    return "UNKNOWN"

def group_by_bse(image_paths):
    groups = []
    current = {"BSE": None, "Spectra": [], "Maps": []}
    for img in image_paths:
        typ = classify(img)
        if typ == "BSE":
            if current["BSE"]:
                groups.append(current)
                current = {"BSE": None, "Spectra": [], "Maps": []}
            current["BSE"] = img
        elif typ == "SPEC":
            current["Spectra"].append(img)
        elif typ == "MAP":
            current["Maps"].append(img)
    if current["BSE"]:
        groups.append(current)
    return groups

def build_doc(groups, output_file):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    for idx, g in enumerate(groups, 1):
        doc.add_page_break()

        # Header
        p = doc.add_paragraph()
        p.alignment = 1  # Centered
        run1 = p.add_run(f"Feature ID-{idx:04d}:")
        run1.bold = True
        run1.underline = True
        run1.font.size = Pt(18)
        run2 = p.add_run(" XXXX")
        run2.bold = True
        run2.font.size = Pt(18)
        p.paragraph_format.space_after = 0

        # BSE
        if g["BSE"]:
            p = doc.add_paragraph()
            p.alignment = 1
            run = p.add_run()
            run.add_picture(str(g["BSE"]), width=Cm(BSE_FMT[1]), height=Cm(BSE_FMT[0]))
            p.paragraph_format.space_after = 0

        specs = g["Spectra"]
        maps = g["Maps"]

        # If exactly 5 spectra: first goes beside BSE
        if len(specs) == 5:
            p = doc.paragraphs[-1]
            run = p.add_run()
            run.add_picture(str(specs[0]), width=Cm(SPEC_FMT[1]), height=Cm(SPEC_FMT[0]))
            specs = specs[1:]

        # Spectra rows (2 per row)
        for i in range(0, len(specs), 2):
            p = doc.add_paragraph()
            p.alignment = 1
            p.paragraph_format.space_after = 0
            for j in range(2):
                if i + j < len(specs):
                    p.add_run().add_picture(str(specs[i + j]), width=Cm(SPEC_FMT[1]), height=Cm(SPEC_FMT[0]))

        # If maps need to go on new page
        if (len(specs) > 6) or (len(specs) > 5 and len(maps) > 4) or (len(maps) > 8):
            doc.add_page_break()

        # Maps (4 per row)
        for i in range(0, len(maps), 4):
            p = doc.add_paragraph()
            p.alignment = 1
            p.paragraph_format.space_after = 0
            for j in range(4):
                if i + j < len(maps):
                    p.add_run().add_picture(str(maps[i + j]), width=Cm(MAP_FMT[1]), height=Cm(MAP_FMT[0]))

    doc.save(output_file)
    print(f"✅ Saved: {output_file}")

def main():
    root = Tk()
    root.withdraw()
    input_path = filedialog.askopenfilename(title="Select your exported EDX Word file", filetypes=[("Word Documents", "*.docx")])
    if not input_path:
        print("❌ No file selected.")
        return

    print("📂 Reading:", input_path)
    tmp = Path(tempfile.mkdtemp())
    ordered_images = extract_ordered_images(input_path, tmp)
    groups = group_by_bse(ordered_images)

    input_path = Path(input_path)
    output_name = f"Modified {input_path.stem}.docx"
    output_path = input_path.with_name(output_name)

    build_doc(groups, output_path)

if __name__ == "__main__":
    main()
