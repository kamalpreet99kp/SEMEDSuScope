import os
import zipfile
import tempfile
from pathlib import Path
from PIL import Image, ImageOps
from tkinter import Tk, filedialog
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# === Original EDX image sizes (cm) ===
BSE_ORIG = (12.16, 15.66)
SPEC_ORIG = (9.00, 15.37)
MAP_ORIG = (7.02, 7.21)
TOL = 0.2

# === Formatted target sizes (cm) ===git remote -v

BSE_FMT = (5.34, 7.25)
SPEC_FMT = (4.75, 7.64)
MAP_FMT = (3.76, 3.66)

# Microscope (uScope) image size: 90% of BSE height
MIC_SCALE = 0.95
MIC_HEIGHT = BSE_FMT[0] * MIC_SCALE  # we will set only height to keep aspect ratio

# White patch at bottom of microscope images (fraction of image height)
PATCH_RATIO = 0.08  # 5% extra white strip at the bottom


def extract_ordered_images(docx_file, extract_dir):
    extract_dir = Path(extract_dir)
    extract_dir.mkdir(parents=True, exist_ok=True)
    doc = Document(docx_file)
    rels = doc.part._rels
    img_order = []

    for p in doc.paragraphs:
        for run in p.runs:
            if "graphic" in run._element.xml and "r:embed=" in run._element.xml:
                rId = run._element.xpath(".//a:blip")[0].get(qn("r:embed"))
                if rId in rels:
                    target = rels[rId]._target
                    if hasattr(target, 'partname') and "media" in target.partname:
                        fname = Path(target.partname).name
                        img_order.append(fname)

    image_map = {}
    with zipfile.ZipFile(docx_file, 'r') as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                img_data = z.read(name)
                fname = name.split("/")[-1]
                out_path = extract_dir / fname
                with open(out_path, "wb") as f:
                    f.write(img_data)
                image_map[fname] = out_path

    ordered_paths = [image_map[f] for f in img_order if f in image_map]
    return ordered_paths


def size_cm(img_path):
    img = Image.open(img_path)
    dpi = img.info.get("dpi", (96, 96))[0]
    w, h = img.size
    return round(h / dpi * 2.54, 2), round(w / dpi * 2.54, 2)


def classify(img_path):
    h, w = size_cm(img_path)
    if abs(h - BSE_ORIG[0]) < TOL and abs(w - BSE_ORIG[1]) < TOL:
        return "BSE"
    elif abs(h - SPEC_ORIG[0]) < TOL and abs(w - SPEC_ORIG[1]) < TOL:
        return "SPEC"
    elif abs(h - MAP_ORIG[0]) < TOL and abs(w - MAP_ORIG[1]) < TOL:
        return "MAP"
    return "UNKNOWN"


def group_by_bse(image_paths):
    groups = []
    current = {"BSE": None, "Spectra": [], "Maps": []}
    for img in image_paths:
        typ = classify(img)
        if typ == "BSE":
            if current["BSE"]:
                groups.append(current)
                current = {"BSE": None, "Spectra": [], "Maps": []}
            current["BSE"] = img
        elif typ == "SPEC":
            current["Spectra"].append(img)
        elif typ == "MAP":
            current["Maps"].append(img)
    if current["BSE"]:
        groups.append(current)
    return groups


def load_microscope_images(cropout_dir):
    cropout_dir = Path(cropout_dir)
    if not cropout_dir.is_dir():
        return []

    exts = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}
    imgs = [p for p in cropout_dir.iterdir() if p.is_file() and p.suffix.lower() in exts]
    imgs.sort(key=lambda p: p.name)
    return imgs


def prepare_microscope_image(img_path, tmp_dir):
    """
    Apply EXIF-based orientation and add a white patch at the bottom
    (5% of image height) so the visible content appears slightly raised.
    """
    img_path = Path(img_path)
    tmp_dir = Path(tmp_dir) if tmp_dir is not None else img_path.parent

    with Image.open(img_path) as im:
        # Correct orientation from EXIF (e.g., Windows Rotate Right)
        im_upright = ImageOps.exif_transpose(im)

        # Ensure RGB (avoid issues with alpha when pasting on white)
        if im_upright.mode not in ("RGB", "L"):
            im_upright = im_upright.convert("RGB")

        w, h = im_upright.size
        patch_h = int(h * PATCH_RATIO)

        # New image with extra white strip at the bottom
        new_h = h + patch_h
        white_bg = Image.new("RGB", (w, new_h), (255, 255, 255))
        white_bg.paste(im_upright, (0, 0))

        out_path = tmp_dir / (img_path.stem + "_upright_padded" + img_path.suffix)
        white_bg.save(out_path)

    return out_path


def build_doc(groups, output_file, microscope_images=None, tmp_dir=None):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    if microscope_images is None:
        microscope_images = []

    for idx, g in enumerate(groups, 1):
        doc.add_page_break()

        # Header
        p = doc.add_paragraph()
        p.alignment = 1
        run = p.add_run(f"Grain No. {idx}")
        run.bold = True
        run.font.size = Pt(18)
        p.paragraph_format.space_after = 0

        # Microscope image for this grain (if available)
        mic_img = microscope_images[idx - 1] if idx - 1 < len(microscope_images) else None

        # BSE + microscope image row
        if g["BSE"]:
            p = doc.add_paragraph()
            p.alignment = 1
            p.paragraph_format.space_after = 0

            # BSE (left) – unchanged
            run = p.add_run()
            run.add_picture(str(g["BSE"]), width=Cm(BSE_FMT[1]), height=Cm(BSE_FMT[0]))

            # Microscope (right) – padded at bottom, 90% height, aspect ratio preserved
            if mic_img is not None:
                mic_prepared = prepare_microscope_image(mic_img, tmp_dir)
                run2 = p.add_run()
                # Only set height to keep aspect ratio correct
                run2.add_picture(str(mic_prepared), height=Cm(MIC_HEIGHT))

        specs = list(g["Spectra"])
        maps = g["Maps"]

        # Spectra rows (2 per row)
        for i in range(0, len(specs), 2):
            p = doc.add_paragraph()
            p.alignment = 1
            p.paragraph_format.space_after = 0
            for j in range(2):
                if i + j < len(specs):
                    p.add_run().add_picture(
                        str(specs[i + j]),
                        width=Cm(SPEC_FMT[1]),
                        height=Cm(SPEC_FMT[0])
                    )

        # If maps need to go on new page
        if (len(specs) > 6) or (len(specs) > 5 and len(maps) > 4) or (len(maps) > 8):
            doc.add_page_break()

        # Maps (4 per row)
        for i in range(0, len(maps), 4):
            p = doc.add_paragraph()
            p.alignment = 1
            p.paragraph_format.space_after = 0
            for j in range(4):
                if i + j < len(maps):
                    p.add_run().add_picture(
                        str(maps[i + j]),
                        width=Cm(MAP_FMT[1]),
                        height=Cm(MAP_FMT[0])
                    )

    doc.save(output_file)
    print(f"✅ Saved: {output_file}")


def main():
    root = Tk()
    root.withdraw()
    input_path = filedialog.askopenfilename(
        title="Select your exported EDX Word file",
        filetypes=[("Word Documents", "*.docx")]
    )
    if not input_path:
        print("❌ No file selected.")
        return

    input_path = Path(input_path)
    print("📂 Reading:", input_path)

    base_dir = input_path.parent
    cropout_dir = base_dir / "cropout"
    microscope_images = []

    if cropout_dir.is_dir():
        microscope_images = load_microscope_images(cropout_dir)
        print(f"🔍 Found {len(microscope_images)} microscope image(s) in {cropout_dir}")
    else:
        print("⚠️ No 'cropout' folder found next to the Word file. Microscope images will be skipped.")

    tmp = Path(tempfile.mkdtemp())
    ordered_images = extract_ordered_images(input_path, tmp)
    groups = group_by_bse(ordered_images)

    output_name = f"Modified {input_path.stem}.docx"
    output_path = input_path.with_name(output_name)

    build_doc(groups, output_path, microscope_images=microscope_images, tmp_dir=tmp)


if __name__ == "__main__":
    main()
