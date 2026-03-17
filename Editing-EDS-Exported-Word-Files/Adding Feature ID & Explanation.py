import pandas as pd
from pathlib import Path
from docx import Document
from tkinter import Tk, filedialog

def update_headers_from_excel_single_prompt():
    # === Step 1: Ask for BOTH Excel and Word files ===
    root = Tk()
    root.withdraw()
    selected_files = filedialog.askopenfilenames(
        title="Select BOTH the Excel file and the Word file",
        filetypes=[("Word/Excel files", "*.docx *.xlsx")]
    )
    if not selected_files or len(selected_files) < 2:
        print("❌ You must select both a Word and an Excel file.")
        return

    # === Step 2: Identify which file is Word and which is Excel ===
    word_path = None
    excel_path = None
    for f in selected_files:
        if f.lower().endswith(".docx"):
            word_path = Path(f)
        elif f.lower().endswith(".xlsx"):
            excel_path = Path(f)

    if not word_path or not excel_path:
        print("❌ Could not find both a Word (.docx) and Excel (.xlsx) file.")
        return

    # === Step 3: Load Excel Data ===
    raw_df = pd.read_excel(excel_path, sheet_name="Sheet2", usecols=[0, 2], header=None, engine="openpyxl")
    first_row = list(raw_df.iloc[0])
    if all(str(x).lower() in ["id", "type"] for x in first_row):
        df = raw_df.iloc[1:].copy()
    else:
        df = raw_df.copy()

    df.columns = ["Id", "Type"]
    df.dropna(subset=["Id", "Type"], inplace=True)

    # === Step 4: Load Word file and process headers ===
    doc = Document(word_path)
    headers = [p for p in doc.paragraphs if "Feature ID-" in p.text and "XXXX" in p.text]

    for i, p in enumerate(headers):
        if i >= len(df):
            break

        id_str = str(df.iloc[i]["Id"]).replace("*", "").strip()
        type_str = str(df.iloc[i]["Type"]).replace("*", "").strip()

        # Clear existing runs
        for r in p.runs:
            p._element.remove(r._element)

        # Add formatted replacements
        run1 = p.add_run(f"Feature ID-{id_str}:")
        run1.bold = True
        run1.underline = True
        run1.font.size = p.style.font.size

        run2 = p.add_run(f" {type_str}")
        run2.bold = True
        run2.font.size = p.style.font.size

    # === Step 5: Save updated Word file ===
    new_name = word_path.with_name(f"Updated {word_path.stem}.docx")
    doc.save(new_name)
    print(f"✅ Done! Saved as: {new_name.name}")

if __name__ == "__main__":
    update_headers_from_excel_single_prompt()
